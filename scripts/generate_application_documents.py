from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
PUBLIC = ROOT / "public"

PRESETS = {
    "access_bank_software_engineer": {
        "resume_md": PUBLIC / "resume_access_bank_software_engineer.md",
        "cover_md": PUBLIC / "cover_letter_access_bank_software_engineer.md",
        "resume_docx": PUBLIC / "Olorunfemi_John_Access_Bank_Software_Engineer_Resume.docx",
        "resume_pdf": PUBLIC / "Olorunfemi_John_Access_Bank_Software_Engineer_Resume.pdf",
        "cover_docx": PUBLIC / "Olorunfemi_John_Access_Bank_Software_Engineer_Cover_Letter.docx",
        "cover_pdf": PUBLIC / "Olorunfemi_John_Access_Bank_Software_Engineer_Cover_Letter.pdf",
        "resume_title": "Olorunfemi John - Access Bank Software Engineer Resume",
        "cover_title": "Olorunfemi John - Access Bank Software Engineer Cover Letter",
    },
    "myhabitra_mobile_engineer": {
        "resume_md": PUBLIC / "resume_myhabitra_mobile_engineer.md",
        "cover_md": PUBLIC / "cover_letter_myhabitra_mobile_engineer.md",
        "resume_docx": PUBLIC / "Olorunfemi_John_MyHabitra_Mobile_Engineer_Resume.docx",
        "resume_pdf": PUBLIC / "Olorunfemi_John_MyHabitra_Mobile_Engineer_Resume.pdf",
        "cover_docx": PUBLIC / "Olorunfemi_John_MyHabitra_Mobile_Engineer_Cover_Letter.docx",
        "cover_pdf": PUBLIC / "Olorunfemi_John_MyHabitra_Mobile_Engineer_Cover_Letter.pdf",
        "resume_title": "Olorunfemi John - MyHabitra Mobile Engineer Resume",
        "cover_title": "Olorunfemi John - MyHabitra Mobile Engineer Cover Letter",
    },
}


LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ACCENT = RGBColor(27, 54, 93)
ACCENT_HEX = "#1B365D"


def normalize_text(text: str) -> str:
    text = LINK_RE.sub(r"\1: \2", text)
    text = text.replace("**", "")
    return text.strip()


def add_page_number(run_paragraph) -> None:
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), "PAGE")
    run_paragraph._element.append(fld_simple)


def set_document_margins(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def add_footer(document: Document) -> None:
    footer = document.sections[0].footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    add_page_number(para)


def set_paragraph_border(paragraph, position: str = "bottom", color: str = "1B365D", size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{position}"))
    if border is None:
        border = OxmlElement(f"w:{position}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), color)


def add_markdown_runs(paragraph, text: str, font_name: str = "Calibri", font_size: float = 10.5) -> None:
    last = 0
    for match in BOLD_RE.finditer(text):
        if match.start() > last:
            run = paragraph.add_run(normalize_text(text[last:match.start()]))
            run.font.name = font_name
            run.font.size = Pt(font_size)
        run = paragraph.add_run(normalize_text(match.group(1)))
        run.bold = True
        run.font.name = font_name
        run.font.size = Pt(font_size)
        last = match.end()
    if last < len(text):
        run = paragraph.add_run(normalize_text(text[last:]))
        run.font.name = font_name
        run.font.size = Pt(font_size)


def parse_resume(md_path: Path) -> dict:
    lines = [line.rstrip() for line in md_path.read_text(encoding="utf-8").splitlines()]
    meaningful = [line for line in lines if line.strip() and line.strip() != "---"]

    name = normalize_text(meaningful[0].replace("# ", "", 1))
    title = normalize_text(meaningful[1])
    contact = normalize_text(meaningful[2])
    links = normalize_text(meaningful[3])

    sections = []
    current_section = None
    current_item = None

    for raw in meaningful[4:]:
        if raw.startswith("## "):
            current_section = {"title": raw[3:].strip(), "items": []}
            sections.append(current_section)
            current_item = None
        elif raw.startswith("### "):
            current_item = {"heading": raw[4:].strip(), "lines": []}
            if current_section is None:
                current_section = {"title": "", "items": []}
                sections.append(current_section)
            current_section["items"].append(current_item)
        else:
            if current_section is None:
                continue
            if current_item is None:
                current_item = {"heading": None, "lines": []}
                current_section["items"].append(current_item)
            current_item["lines"].append(raw)

    return {
        "name": name,
        "title": title,
        "contact": contact,
        "links": links,
        "sections": sections,
    }


def build_resume_docx(data: dict, output_path: Path) -> None:
    doc = Document()
    set_document_margins(doc)
    add_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(data["name"])
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(18)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(data["title"])
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor(68, 68, 68)

    for header_line in (data["contact"], data["links"]):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(header_line)
        r.font.name = "Calibri"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(55, 55, 55)

    divider = doc.add_paragraph()
    divider.paragraph_format.space_before = Pt(2)
    divider.paragraph_format.space_after = Pt(7)
    set_paragraph_border(divider, color="1B365D", size="8")

    for section in data["sections"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(section["title"].upper())
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(11.5)
        r.font.color.rgb = ACCENT
        set_paragraph_border(p, color="1B365D", size="6")

        for item in section["items"]:
            if item["heading"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(item["heading"])
                r.bold = True
                r.font.name = "Calibri"
                r.font.size = Pt(10.75)
                r.font.color.rgb = RGBColor(32, 32, 32)

            for line in item["lines"]:
                if line.startswith("- "):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.18)
                    p.paragraph_format.first_line_indent = Inches(-0.18)
                    p.paragraph_format.space_after = Pt(1.5)
                    p.paragraph_format.line_spacing = 1.08
                    lead = p.add_run("- ")
                    lead.font.name = "Calibri"
                    lead.font.size = Pt(10.25)
                    add_markdown_runs(p, line[2:], font_size=10.25)
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(1.5)
                    p.paragraph_format.line_spacing = 1.06
                    if line.startswith("**Live:**") or line.startswith("**GitHub:**") or line.startswith("**Live:**") or line.startswith("**GitHub:**") or " | **GitHub:**" in line:
                        add_markdown_runs(p, line, font_size=9.8)
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(85, 85, 85)
                    else:
                        add_markdown_runs(p, line, font_size=10.25)

    temp_path = output_path.with_suffix(output_path.suffix + ".tmp")
    doc.save(temp_path)
    temp_path.replace(output_path)


def build_resume_pdf(data: dict, output_path: Path, document_title: str) -> None:
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle(
        "ResumeName",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        textColor=colors.HexColor(ACCENT_HEX),
        fontSize=18,
        leading=21,
        alignment=TA_CENTER,
        spaceAfter=2,
    )
    title_style = ParagraphStyle(
        "ResumeTitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=11.5,
        leading=13,
        textColor=colors.HexColor("#444444"),
        alignment=TA_CENTER,
        spaceAfter=4,
    )
    header_style = ParagraphStyle(
        "ResumeHeader",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        textColor=colors.HexColor("#333333"),
        leading=11,
        alignment=TA_CENTER,
        spaceAfter=1,
    )
    section_style = ParagraphStyle(
        "ResumeSection",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        textColor=colors.HexColor(ACCENT_HEX),
        fontSize=11.5,
        leading=14,
        spaceBefore=7,
        spaceAfter=3,
        alignment=TA_LEFT,
    )
    heading_style = ParagraphStyle(
        "ResumeHeading",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10.75,
        leading=13,
        spaceBefore=4,
        spaceAfter=2,
        textColor=colors.HexColor("#222222"),
    )
    meta_style = ParagraphStyle(
        "ResumeMeta",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.8,
        textColor=colors.HexColor("#555555"),
        leading=12,
        spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10.25,
        leading=13,
        alignment=TA_JUSTIFY,
        spaceAfter=2,
    )
    bullet_style = ParagraphStyle(
        "ResumeBullet",
        parent=body_style,
        leftIndent=11,
        firstLineIndent=0,
        spaceAfter=2,
    )

    story = [
        Paragraph(data["name"], name_style),
        Paragraph(data["title"], title_style),
        Paragraph(data["contact"], header_style),
        Paragraph(data["links"], header_style),
        Spacer(1, 0.02 * inch),
        HRFlowable(width="100%", thickness=1.1, color=colors.HexColor(ACCENT_HEX), spaceBefore=0, spaceAfter=8),
    ]

    for section in data["sections"]:
        story.append(Paragraph(section["title"].upper(), section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=colors.HexColor(ACCENT_HEX), spaceBefore=0, spaceAfter=4))
        for item in section["items"]:
            if item["heading"]:
                story.append(Paragraph(item["heading"], heading_style))
            for line in item["lines"]:
                if line.startswith("- "):
                    story.append(Paragraph(f"- {normalize_text(line[2:])}", bullet_style))
                elif line.startswith("**Live:**") or line.startswith("**GitHub:**") or " | **GitHub:**" in line:
                    story.append(Paragraph(normalize_text(line), meta_style))
                else:
                    story.append(Paragraph(normalize_text(line), body_style))

    doc = SimpleDocTemplate(
        str(output_path.with_suffix(output_path.suffix + ".tmp")),
        pagesize=LETTER,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
        title=document_title,
        author="Olorunfemi John",
    )
    doc.build(story)
    temp_path = output_path.with_suffix(output_path.suffix + ".tmp")
    temp_path.replace(output_path)


def parse_cover_letter(md_path: Path) -> list[str]:
    lines = [line.strip() for line in md_path.read_text(encoding="utf-8").splitlines()]
    return [normalize_text(line) for line in lines if line]


def build_letter_docx(paragraphs: list[str], output_path: Path) -> None:
    doc = Document()
    set_document_margins(doc)
    add_footer(doc)

    for index, text in enumerate(paragraphs):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.08
        add_markdown_runs(p, text, font_size=10.75)
        if index == 0:
            p.paragraph_format.space_before = Pt(2)

    temp_path = output_path.with_suffix(output_path.suffix + ".tmp")
    doc.save(temp_path)
    temp_path.replace(output_path)


def build_letter_pdf(paragraphs: list[str], output_path: Path, document_title: str) -> None:
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "LetterBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10.75,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=10,
    )
    story = [Paragraph(text, body_style) for text in paragraphs]

    doc = SimpleDocTemplate(
        str(output_path.with_suffix(output_path.suffix + ".tmp")),
        pagesize=LETTER,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.7 * inch,
        title=document_title,
        author="Olorunfemi John",
    )
    doc.build(story)
    temp_path = output_path.with_suffix(output_path.suffix + ".tmp")
    temp_path.replace(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate application documents from markdown sources.")
    parser.add_argument(
        "--preset",
        default="access_bank_software_engineer",
        choices=sorted(PRESETS),
        help="Document preset to build.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    preset = PRESETS[args.preset]

    resume_data = parse_resume(preset["resume_md"])
    cover_paragraphs = parse_cover_letter(preset["cover_md"])

    build_resume_docx(resume_data, preset["resume_docx"])
    build_resume_pdf(resume_data, preset["resume_pdf"], preset["resume_title"])
    build_letter_docx(cover_paragraphs, preset["cover_docx"])
    build_letter_pdf(cover_paragraphs, preset["cover_pdf"], preset["cover_title"])

    print(f"Generated: {preset['resume_docx']}")
    print(f"Generated: {preset['resume_pdf']}")
    print(f"Generated: {preset['cover_docx']}")
    print(f"Generated: {preset['cover_pdf']}")


if __name__ == "__main__":
    main()
